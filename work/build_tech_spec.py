from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\Users\user\Documents\Codex\2026-07-26\ai-jira-confluence\outputs\technical_spec_wms_ai_assistant_v1.2.docx")
BLUE = "2E74B5"
DARK = "1F4D78"
INK = "172033"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"

def set_font(run, size=11, bold=None, color=INK, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = tcPr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar"); tcPr.append(mar)
    for side, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}"); mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")

def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tcW = cell._tc.tcPr.first_child_found_in("w:tcW")
            tcW.set(qn("w:w"), str(width)); tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

def add_para(doc, text="", style=None, bold_label=None):
    p = doc.add_paragraph(style=style)
    if bold_label and text.startswith(bold_label):
        r = p.add_run(bold_label); set_font(r, bold=True)
        r = p.add_run(text[len(bold_label):]); set_font(r)
    else:
        r = p.add_run(text); set_font(r)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text))
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    set_font(p.add_run(text))
    return p

def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0,0); shade(cell, "172033")
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code); run.font.name = "Consolas"; run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas"); run.font.size = Pt(9); run.font.color.rgb = RGBColor(232,239,255)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1); set_table_widths(table, [9360])
    cell = table.cell(0,0); shade(cell, CALLOUT)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(title), bold=True, color=DARK)
    p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(0); set_font(p.add_run(text))

doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
section.header_distance = Inches(.492); section.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK,8,4)]:
    st = styles[name]; st.font.name = "Calibri"; st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color); st.font.bold = True
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("WMS AI Assistant | Техническая спецификация"), size=9, color="6B7280")
footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("Внутренний документ | Демонстрационный прототип"), size=9, color="6B7280")

# Masthead
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
set_font(p.add_run("ТЕХНИЧЕСКАЯ СПЕЦИФИКАЦИЯ"), size=23, bold=True, color="000000")
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
set_font(p.add_run("Локальный AI-помощник для поддержки WMS: Jira, Confluence и Ollama"), size=14, color="4B5563")
meta = doc.add_table(rows=3, cols=2); set_table_widths(meta,[1800,7560])
for row, (label, value) in zip(meta.rows, [("Версия", "1.2"), ("Дата", "27 июля 2026"), ("Статус", "Тестовый прототип")]):
    shade(row.cells[0], LIGHT); set_font(row.cells[0].paragraphs[0].add_run(label), bold=True)
    set_font(row.cells[1].paragraphs[0].add_run(value))

add_callout(doc, "Цель", "Помочь консультантам первой и второй линии быстро определить первичные проверки, вероятную техническую причину, прошлое решение аналогичного обращения, связанную документацию и задачу разработки.")

doc.add_heading("1. Назначение и границы", level=1)
add_para(doc, "Приложение анализирует обращения пользователей складской WMS. Оно ищет похожие обращения, задачи разработки и документацию, после чего формирует единый безопасный сценарий решения.")
add_bullet(doc, "Поддерживает приёмку, отгрузку, инвентаризацию, ТСД и печать Zebra.")
add_bullet(doc, "Не меняет остатки, права, документы или статусы задач.")
add_bullet(doc, "Рекомендации требуют проверки консультантом до выполнения действий в продуктивной среде.")

doc.add_heading("2. Пользователи и роли", level=1)
roles = doc.add_table(rows=1, cols=3); set_table_widths(roles,[1800,3300,4260])
for cell, text in zip(roles.rows[0].cells,["Роль","Задача","Результат"]): shade(cell,LIGHT); set_font(cell.paragraphs[0].add_run(text),bold=True)
for role, task, result in [
    ("Кладовщик","Создаёт обращение и отвечает на уточняющие вопросы.","Передаёт понятные наблюдения без технических терминов."),
    ("Консультант 1-2 линии","Проверяет первичные причины, выполняет регламент и при необходимости эскалирует дефект.","Получает сценарий решения, документацию и прошлый опыт."),
    ("Разработчик","Разбирает подтверждённый дефект.","Получает ссылку на задачу PROB и контекст ошибки."),
    ("Администратор","Настраивает доступы и модели.","Контролирует интеграции и безопасность.")]:
    cells=roles.add_row().cells
    for cell,text in zip(cells,[role,task,result]): set_font(cell.paragraphs[0].add_run(text))

doc.add_heading("3. Архитектура", level=1)
add_para(doc, "Контекст обращения проходит через слой семантического поиска и три логические AI-роли.")
add_code(doc, "Новое обращение в Jira\n  ├─ Сабагент обращений → похожие WMS-обращения и PROB-задачи\n  ├─ Сабагент документации → страницы Confluence и точки ошибок\n  └─ Агент-оркестратор → готовый сценарий для консультанта")
add_para(doc, "Два сабагента запускаются параллельно. Оркестратор получает только их структурированные выводы, а не полную базу данных.")

doc.add_heading("4. Локальные модели Ollama", level=1)
models = doc.add_table(rows=1, cols=4); set_table_widths(models,[2100,2100,2100,3060])
for cell,text in zip(models.rows[0].cells,["Роль","Модель","API","Назначение"]): shade(cell,LIGHT); set_font(cell.paragraphs[0].add_run(text),bold=True)
for row in [("Семантический поиск","embeddinggemma","/api/embed","Векторы обращений и документации."),("Сабагенты и оркестратор","qwen3.6:latest","/api/chat","Анализ текста и формирование ответа.")]:
    cells=models.add_row().cells
    for cell,text in zip(cells,row): set_font(cell.paragraphs[0].add_run(text))
add_callout(doc, "Важно", "embeddinggemma применяется только для embedding-поиска и не должна использоваться как чат-модель.")

doc.add_heading("5. Процесс анализа обращения", level=1)
for step in [
    "Получить текст обращения, переписку и приложенные данные.",
    "Найти близкие по смыслу WMS-обращения и PROB-задачи через embeddinggemma.",
    "Найти релевантные страницы Confluence, включая точки генерации ошибок.",
    "Передать историю обращений сабагенту обращений, а документацию - сабагенту документации.",
    "Передать оба вывода оркестратору и вернуть сценарий решения."
]: add_number(doc, step)

doc.add_heading("6. Формат результата оркестратора", level=1)
add_code(doc, '''{
  "summary": "Краткое описание ситуации",
  "possible_causes": [{"cause":"...","evidence":"...","confidence":"medium"}],
  "error_locations": [{"error":"...","location":"функция → условие","condition":"..."}],
  "solution_scenario": {"goal":"...","steps":["..."],"verification":["..."],"escalation":["..."]},
  "questions": ["..."],
  "risk_note": "..."
}''')

doc.add_heading("7. Демонстрационные данные", level=1)
add_bullet(doc, "Всего 530 обезличенных тестовых заявок.")
add_bullet(doc, "500 WMS-обращений с ключами WMS-*.")
add_bullet(doc, "30 задач разработки типа Development Bug с ключами PROB-*.")
add_bullet(doc, "175 WMS-кейсов содержат переписку консультанта с пользователем.")
add_bullet(doc, "250 WMS-кейсов являются переформулированными вариантами предыдущих сценариев и расширяют семантический поиск.")
add_bullet(doc, "20 страниц демо-Confluence в пространстве WMSDEV.")

doc.add_heading("8. Поиск места возникновения ошибки", level=1)
add_para(doc, "Страницы разработки содержат примерную логику, код ошибки, условие и функцию, в которой ошибка должна возникать. Поиск повышает релевантность страницы при совпадении HTTP-кода или идентификатора ошибки.")
example = doc.add_table(rows=1, cols=3); set_table_widths(example,[2600,3600,3160])
for cell,text in zip(example.rows[0].cells,["Ошибка","Место в логике","Условие"]): shade(cell,LIGHT); set_font(cell.paragraphs[0].add_run(text),bold=True)
cells=example.add_row().cells
for cell,text in zip(cells,["409 INVENTORY_CELL_NOT_LOCKED","postCount() → isCellLockedBy()","Ячейка разблокирована или заблокирована другим листом."]): set_font(cell.paragraphs[0].add_run(text))

doc.add_heading("9. Автоматическая обработка обращений в Jira", level=1)
add_bullet(doc, "Webhook Jira запускает анализ, когда в выбранных проектах появляется новое обращение.")
add_bullet(doc, "Сервисная учётная запись должна видеть обращения в этих проектах и иметь право добавлять комментарии.")
add_bullet(doc, "После анализа помощник добавляет один комментарий: что проверить в первую очередь, какое решение помогло в похожем случае, ссылки на документацию и связанную PROB-задачу.")
add_bullet(doc, "Перед публикацией система проверяет, не был ли такой комментарий уже добавлен, чтобы не создавать дубли.")

doc.add_heading("10. Интеграция Confluence", level=1)
add_bullet(doc, "Используются URL Confluence, e-mail сервисной учётной записи, API token и ключи разрешённых пространств.")
add_bullet(doc, "Поиск ограничивается страницами, к которым сервисная учётная запись имеет доступ.")
add_bullet(doc, "Найденные страницы и точки ошибок передаются сабагенту документации.")

doc.add_heading("11. Безопасность и ограничения", level=1)
add_bullet(doc, "Локальный Ollama не отправляет текст обращений в OpenAI.")
add_bullet(doc, "Текущая чат-модель анализирует текст. Для скриншотов требуется отдельная локальная vision-модель.")
add_bullet(doc, "Комментарии в Jira публикует только сервисная учётная запись; пользовательский доступ к чужим обращениям не используется.")

doc.add_heading("12. Критерии готовности", level=1)
for item in [
    "Новое обращение автоматически поступает в анализ.",
    "Находятся похожие WMS-обращения, PROB-задачи и страницы Confluence.",
    "В ответе есть безопасный сценарий решения и условия эскалации.",
    "При известной ошибке отображается место её возникновения в логике.",
    "Для каждого нового обращения в Jira публикуется один понятный комментарий без повторов.",
    "При недоступности сабагента используется резервный вывод со статусом ошибки."
]: add_bullet(doc, item)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "Техническая спецификация WMS AI-помощника"
doc.core_properties.subject = "Jira, Confluence и Ollama"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT)
